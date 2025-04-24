import streamlit as st
import pdfplumber
import openai
import pandas as pd
from langchain.document_loaders import PyPDFLoader
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from fpdf import FPDF
import os
import re
import time
import matplotlib.pyplot as plt
import docx
from docx import Document
from io import BytesIO
import base64
from datetime import datetime

OPENAI_API_KEY = ""
openai.api_key = OPENAI_API_KEY

# Initialize session state variables
if "screening_done" not in st.session_state:
    st.session_state.screening_done = False
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "resume_history" not in st.session_state:
    st.session_state.resume_history = []
if "comparison_results" not in st.session_state:
    st.session_state.comparison_results = []
if "current_resume_text" not in st.session_state:
    st.session_state.current_resume_text = ""
if "current_job_description" not in st.session_state:
    st.session_state.current_job_description = ""
if "ats_simulation_results" not in st.session_state:
    st.session_state.ats_simulation_results = None
if "keywords_extracted" not in st.session_state:
    st.session_state.keywords_extracted = []

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    try:
        loader = PyPDFLoader(pdf_path)
        pages = loader.load()
        text = "\n".join([page.page_content for page in pages])
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

# Function to extract text from uploaded file (PDF or DOCX)
def extract_text_from_file(uploaded_file):
    # Save the uploaded file temporarily
    temp_file_path = f"temp_file_{int(time.time())}"
    with open(temp_file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    if uploaded_file.name.endswith('.pdf'):
        text = extract_text_from_pdf(temp_file_path)
    elif uploaded_file.name.endswith('.docx'):
        try:
            doc = Document(temp_file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            st.error(f"Error extracting text from DOCX: {e}")
            text = ""
    else:
        text = ""
    
    # Clean up the temporary file
    try:
        os.remove(temp_file_path)
    except:
        pass
    
    return text

# Function to extract required skills from the job description
def extract_required_skills(job_description):
    prompt = f"""
    Given the following job description:
    {job_description}
    
    Extract a list of key required skills, qualifications, and experiences for the role.
    For each item, include a brief explanation of why it's important.
    Format as a JSON list with "skill" and "importance" keys.
    """
    
    llm = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY)
    response = llm.predict(prompt)
    
    # Extract just the skills as a list for simpler display
    skills_list = []
    try:
        # Try to parse as JSON if it's formatted correctly
        import json
        skills_data = json.loads(response)
        skills_list = [item["skill"] for item in skills_data]
    except:
        # Fallback to simple parsing
        skills_list = [skill.strip() for skill in response.split(",")]
    
    return skills_list

# Function to extract keywords from job description
def extract_job_keywords(job_description):
    prompt = f"""
    Analyze this job description and extract the top 10-15 keywords that an ATS (Applicant Tracking System) would likely scan for:
    
    {job_description}
    
    Return only a comma-separated list of keywords.
    """
    
    llm = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY)
    response = llm.predict(prompt)
    keywords = [keyword.strip() for keyword in response.split(",")]
    return keywords

# Function to process resume with OpenAI
def analyze_resume_with_openai(text, job_description):
    st.session_state.keywords_extracted = extract_job_keywords(job_description)
    
    prompt = f"""
    Given the following resume text:
    {text}
    
    And the job description:
    {job_description}
    
    Perform a detailed analysis:
    1. Extract key skills and experiences from the resume.
    2. Match them against the requirements in the job description.
    3. Assign a relevance score (0-100) based on how well the resume matches the job.
    4. Provide a brief summary of strengths.
    5. Identify missing skills or experiences that would improve the match.
    6. Suggest specific improvements to better align with the job description.
    
    Return output in this format:
    - Matched Skills: [List each matched skill]
    - Missing Skills: [List each missing skill]
    - Score: [X]%
    - Strengths: [Brief summary of candidate's strengths]
    - Improvement Areas: [Specific suggestions for improvement]
    - ATS Compatibility: [Assess how well the resume would perform with ATS systems]
    """
    
    response = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY).predict(prompt)
    return response

# Function to generate optimized resume sections
def generate_optimized_sections(resume_text, job_description, section):
    prompt = f"""
    Based on this resume:
    {resume_text}
    
    And this job description:
    {job_description}
    
    Rewrite the {section} section to better align with the job requirements.
    Make it more impactful, keyword-rich, and optimized for ATS systems.
    Include measurable achievements where appropriate.
    
    Return only the optimized {section} section text.
    """
    
    response = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY).predict(prompt)
    return response

# Function to simulate ATS parsing
def simulate_ats_parsing(resume_text, job_keywords):
    # Count keyword occurrences
    keyword_counts = {}
    for keyword in job_keywords:
        count = len(re.findall(r'\b' + re.escape(keyword.lower()) + r'\b', resume_text.lower()))
        keyword_counts[keyword] = count
    
    # Calculate match percentage
    total_keywords = len(job_keywords)
    matched_keywords = sum(1 for count in keyword_counts.values() if count > 0)
    match_percentage = (matched_keywords / total_keywords) * 100 if total_keywords > 0 else 0
    
    # Get ATS format recommendations
    ats_format_prompt = f"""
    Analyze this resume text for ATS compatibility issues:
    
    {resume_text}
    
    Identify formatting problems, missing sections, or other issues that might prevent an ATS from correctly parsing the resume.
    Provide specific formatting recommendations to improve ATS compatibility.
    """
    
    format_recommendations = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY).predict(ats_format_prompt)
    
    return {
        "keyword_counts": keyword_counts,
        "match_percentage": match_percentage,
        "format_recommendations": format_recommendations
    }

# Function to generate a PDF report
def generate_pdf_report(candidate_name, analysis_result, job_description, output_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Add cover page
    pdf.add_page()
    pdf.set_font("Arial", style='B', size=24)
    pdf.cell(200, 20, "Resume Analysis Report", ln=True, align='C')
    pdf.set_font("Arial", style='B', size=16)
    pdf.cell(200, 20, f"Candidate: {candidate_name}", ln=True, align='C')
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True, align='C')
    
    # Add job description summary
    pdf.add_page()
    pdf.set_font("Arial", style='B', size=16)
    pdf.cell(200, 10, "Job Description Summary", ln=True)
    pdf.set_font("Arial", size=12)
    
    # Summarize job description
    job_summary_prompt = f"Summarize this job description in 3-4 bullet points: {job_description}"
    job_summary = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY).predict(job_summary_prompt)
    
    pdf.multi_cell(0, 10, job_summary)
    
    # Add analysis results
    pdf.add_page()
    pdf.set_font("Arial", style='B', size=16)
    pdf.cell(200, 10, "Resume Analysis Results", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, analysis_result)
    
    # Add recommendations page
    pdf.add_page()
    pdf.set_font("Arial", style='B', size=16)
    pdf.cell(200, 10, "Recommendations for Improvement", ln=True)
    
    # Generate specific recommendations
    recommendations_prompt = f"""
    Based on this resume analysis:
    {analysis_result}
    
    Provide 5 specific, actionable recommendations to improve the resume for this job description:
    {job_description}
    
    Format as a numbered list.
    """
    
    recommendations = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY).predict(recommendations_prompt)
    
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, recommendations)
    
    pdf.output(output_path)
    return output_path

# Function to create a comparison chart
def create_comparison_chart(scores):
    labels = [f"Attempt {i+1}" for i in range(len(scores))]
    
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.bar(labels, scores, color='purple')
    ax.set_ylim(0, 100)
    ax.set_ylabel('Match Score (%)')
    ax.set_title('Resume Improvement Progress')
    
    # Add horizontal guidelines
    ax.axhline(y=50, color='r', linestyle='--', alpha=0.3)
    ax.axhline(y=70, color='y', linestyle='--', alpha=0.3)
    ax.axhline(y=90, color='g', linestyle='--', alpha=0.3)
    
    # Save to a temporary buffer
    buf = BytesIO()
    fig.savefig(buf, format="png")
    buf.seek(0)
    
    # Create the base64 string for display
    img_str = base64.b64encode(buf.read()).decode()
    return img_str

# Function to generate DOCX resume
def generate_improved_resume(resume_text, analysis_result, job_description):
    # Extract sections to improve
    sections_prompt = f"""
    Based on this resume:
    {resume_text}
    
    And this job description:
    {job_description}
    
    Create an improved version of the resume with these sections:
    1. Professional Summary (tailored to the job)
    2. Skills (optimized with relevant keywords)
    3. Experience (reformatted with achievements and keywords)
    4. Education
    
    Format each section with clear headings.
    """
    
    improved_resume = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY).predict(sections_prompt)
    
    # Create Word document
    doc = Document()
    doc.add_heading('Optimized Resume', 0)
    
    # Add improved content
    for section in improved_resume.split('\n\n'):
        if section.strip():
            if section.strip().endswith(':'):
                doc.add_heading(section.strip(), level=1)
            else:
                doc.add_paragraph(section.strip())
    
    # Save to a BytesIO object
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

# Custom CSS
st.markdown("""
<style>
    .stApp {
        background-color: #f0f2f6;
    }
    .main-header {
        color: #8A2BE2;
        font-size: 36px;
        font-weight: bold;
        text-align: center;
        margin-bottom: 20px;
        background: linear-gradient(90deg, #9370DB, #8A2BE2);
        color: white;
        padding: 20px;
        border-radius: 10px;
    }
    .section-header {
        background-color: #9370DB;
        color: white;
        padding: 10px;
        border-radius: 5px;
        margin-top: 20px;
        margin-bottom: 20px;
    }
    .info-box {
        background-color: #e6e6fa;
        padding: 15px;
        border-radius: 5px;
        margin-bottom: 15px;
        border-left: 5px solid #8A2BE2;
    }
    .metric-card {
        background-color: white;
        padding: 20px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        text-align: center;
    }
    .keyword-tag {
        display: inline-block;
        background-color: #9370DB;
        color: white;
        padding: 5px 10px;
        margin: 5px;
        border-radius: 15px;
        font-size: 0.8em;
    }
    .custom-table {
        width: 100%;
        text-align: left;
        border-collapse: collapse;
    }
    .custom-table th, .custom-table td {
        padding: 10px;
        border: 1px solid #ddd;
    }
    .custom-table th {
        background-color: #9370DB;
        color: white;
    }
    .custom-table tr:nth-child(even) {
        background-color: #f2f2f2;
    }
    .sidebar-content {
        background-color: #f8f9fa;
        padding: 15px;
        border-radius: 5px;
    }
    .stButton>button {
        background-color: #8A2BE2;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 24px;
        font-weight: bold;
    }
    .stProgress .st-bo {
        background-color: #9370DB;
    }
    /* Chat styling */
    .chat-message {
        padding: 15px;
        border-radius: 10px;
        margin-bottom: 10px;
        display: flex;
    }
    .chat-message.user {
        background-color: #e6e6fa;
    }
    .chat-message.bot {
        background-color: #f0f0f0;
    }
</style>
""", unsafe_allow_html=True)

# Create app layout with multiple tabs
st.markdown("<h1 class='main-header'>🚀 Fix & Fit: AI Resume Optimizer</h1>", unsafe_allow_html=True)

# Sidebar navigation and instructions
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/4661/4661321.png", width=100)
    st.markdown("<div class='sidebar-content'>", unsafe_allow_html=True)
    st.subheader("How It Works")
    st.markdown("""
    1. Upload your resume
    2. Enter the job description
    3. Get AI analysis & suggestions
    4. Optimize your resume
    5. Check ATS compatibility
    """)
    st.markdown("</div>", unsafe_allow_html=True)
    
    if len(st.session_state.resume_history) > 0:
        st.subheader("Resume History")
        for i, entry in enumerate(st.session_state.resume_history):
            st.write(f"{i+1}. {entry['timestamp']} - Score: {entry['score']}%")

# Create tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs(["📝 Resume Analysis", "🔍 ATS Simulator", "✏️ Resume Optimizer", "📊 Progress Tracker", "💬 AI Assistant"])

# TAB 1: RESUME ANALYSIS
with tab1:
    st.markdown("<h3 class='section-header'>Upload Your Resume & Job Description</h3>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        uploaded_file = st.file_uploader("**📂 Upload your Resume**", type=["pdf", "docx"], help="We support PDF and Word formats")
    
    with col2:
        job_description = st.text_area("**📋 Enter the Job Description**", 
                                       st.session_state.current_job_description, 
                                       height=150,
                                       help="Copy and paste the full job posting")
    
    # Save current job description in session state
    if job_description:
        st.session_state.current_job_description = job_description
    
    analyze_button = st.button("🔍 Analyze Resume", use_container_width=True)
    
    if uploaded_file is not None and analyze_button:
        with st.spinner("Analyzing your resume against the job description..."):
            # Extract text from file
            text = extract_text_from_file(uploaded_file)
            st.session_state.current_resume_text = text
            
            # Process the resume
            if text and job_description:
                analysis_result = analyze_resume_with_openai(text, job_description)
                st.session_state.analysis_result = analysis_result
                st.session_state.screening_done = True
                
                # Extract score for history tracking
                try:
                    score_text = analysis_result.split("Score:")[1].split("%")[0].strip()
                    score = int(re.sub(r'[^0-9]', '', score_text))
                except:
                    score = 0
                
                # Add to history
                st.session_state.resume_history.append({
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "filename": uploaded_file.name,
                    "score": score
                })
                
                # Create comparison results if this is the first entry
                if len(st.session_state.comparison_results) == 0:
                    st.session_state.comparison_results.append(score)
                else:
                    st.session_state.comparison_results.append(score)
                
                # Generate ATS simulation results
                st.session_state.ats_simulation_results = simulate_ats_parsing(
                    text, 
                    st.session_state.keywords_extracted
                )
                
                # Generate report
                candidate_name = uploaded_file.name.replace(".pdf", "").replace(".docx", "")
                report_path = generate_pdf_report(
                    candidate_name, 
                    analysis_result, 
                    job_description,
                    "resume_report.pdf"
                )
                
                st.success("Analysis complete!")
                st.experimental_rerun()
    
    if st.session_state.screening_done and hasattr(st.session_state, 'analysis_result'):
        st.markdown("<h3 class='section-header'>Analysis Results</h3>", unsafe_allow_html=True)
        
        # Parse the analysis result
        analysis = st.session_state.analysis_result
        
        try:
            matched_skills = analysis.split("Matched Skills:")[1].split("Missing Skills:")[0].strip()
            missing_skills = analysis.split("Missing Skills:")[1].split("Score:")[0].strip()
            score_section = analysis.split("Score:")[1].split("\n")[0].strip()
            score = int(re.sub(r'[^0-9]', '', score_section))
            
            # Try to extract strengths and improvement areas
            strengths = ""
            improvement_areas = ""
            
            if "Strengths:" in analysis:
                strengths = analysis.split("Strengths:")[1].split("Improvement Areas:" if "Improvement Areas:" in analysis else "ATS Compatibility:")[0].strip()
            
            if "Improvement Areas:" in analysis:
                improvement_areas = analysis.split("Improvement Areas:")[1].split("ATS Compatibility:" if "ATS Compatibility:" in analysis else "\n\n")[0].strip()
            
            ats_compatibility = ""
            if "ATS Compatibility:" in analysis:
                ats_compatibility = analysis.split("ATS Compatibility:")[1].strip()
        except Exception as e:
            st.error(f"Error parsing analysis: {e}")
            matched_skills = missing_skills = ""
            score = 0
            strengths = improvement_areas = ats_compatibility = ""
        
        # Create metrics
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
            st.metric("Match Score", f"{score}%")
            
            # Display progress bar
            progress_color = "green" if score > 70 else "orange" if score > 50 else "red"
            st.progress(score/100)
            
            if score >= 70:
                st.success("Your resume is a good match for this job!")
            elif score >= 50:
                st.warning("Your resume has some relevant skills but needs improvement.")
            else:
                st.error("Your resume needs significant improvements for this job.")
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        with col2:
            st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
            st.subheader("Keywords Found")
            
            # Display a few key found keywords
            if hasattr(st.session_state, 'ats_simulation_results') and st.session_state.ats_simulation_results:
                keyword_counts = st.session_state.ats_simulation_results["keyword_counts"]
                found_keywords = [k for k, v in keyword_counts.items() if v > 0]
                missing_keywords = [k for k, v in keyword_counts.items() if v == 0]
                
                st.markdown(f"<p>✅ Found: {len(found_keywords)}/{len(keyword_counts)}</p>", unsafe_allow_html=True)
                st.markdown(f"<p>❌ Missing: {len(missing_keywords)}/{len(keyword_counts)}</p>", unsafe_allow_html=True)
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        # Display detailed results
        st.markdown("<div class='info-box'>", unsafe_allow_html=True)
        
        tab_results1, tab_results2, tab_results3 = st.tabs(["Match Analysis", "Strengths & Improvements", "Full Report"])
        
        with tab_results1:
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("✅ Matched Skills")
                st.markdown(matched_skills)
            
            with col2:
                st.subheader("❌ Missing Skills")
                st.markdown(missing_skills)
        
        with tab_results2:
            if strengths:
                st.subheader("💪 Strengths")
                st.markdown(strengths)
            
            if improvement_areas:
                st.subheader("🔧 Areas for Improvement")
                st.markdown(improvement_areas)
            
            if ats_compatibility:
                st.subheader("🤖 ATS Compatibility")
                st.markdown(ats_compatibility)
        
        with tab_results3:
            st.markdown(analysis)
        
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Download report button
        with open("resume_report.pdf", "rb") as f:
            st.download_button(
                label="📥 Download Full Analysis Report",
                data=f,
                file_name="Resume_Analysis_Report.pdf",
                mime="application/pdf",
                use_container_width=True
            )

# TAB 2: ATS SIMULATOR
with tab2:
    st.markdown("<h3 class='section-header'>ATS Compatibility Check</h3>", unsafe_allow_html=True)
    
    st.markdown("<div class='info-box'>", unsafe_allow_html=True)
    st.write("""
    Applicant Tracking Systems (ATS) are software tools that companies use to scan and filter resumes. 
    This simulator shows how well your resume performs with typical ATS systems and highlights keywords 
    that may be missing.
    """)
    st.markdown("</div>", unsafe_allow_html=True)
    
    if not st.session_state.ats_simulation_results:
        st.warning("Please analyze a resume in the Resume Analysis tab first.")
    else:
        # Display ATS simulation results
        ats_results = st.session_state.ats_simulation_results
        match_percentage = ats_results["match_percentage"]
        keyword_counts = ats_results["keyword_counts"]
        format_recommendations = ats_results["format_recommendations"]
        
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.markdown("<div class='metric-card'>", unsafe_allow_html=True)
            st.metric("ATS Score", f"{int(match_percentage)}%")
            
            # ATS verdict
            if match_percentage >= 70:
                st.success("Your resume is likely to pass ATS screening!")
            elif match_percentage >= 50:
                st.warning("Your resume may pass ATS screening, but it's borderline.")
            else:
                st.error("Your resume is likely to be filtered out by ATS systems.")
            
            st.markdown("</div>", unsafe_allow_html=True)
        
        with col2:
            # Display keyword analysis
            st.subheader("Keyword Analysis")
            
            # Found vs missing keywords
            found_keywords = {k: v for k, v in keyword_counts.items() if v > 0}
            missing_keywords = {k: v for k, v in keyword_counts.items() if v == 0}
            
            col_a, col_b = st.columns(2)
            
            with col_a:
                st.markdown("**Keywords Found:**")
                for keyword, count in found_keywords.items():
                    st.markdown(f"<span class='keyword-tag'>✅ {keyword} ({count})</span>", unsafe_allow_html=True)
            
            with col_b:
                st.markdown("**Keywords Missing:**")
                for keyword in missing_keywords:
                    st.markdown(f"<span class='keyword-tag' style='background-color:#ff6b6b;'>❌ {keyword}</span>", unsafe_allow_html=True)
        
        # Display keyword density visualization
        if found_keywords:
            st.subheader("Keyword Density")
            
            keywords = list(found_keywords.keys())
            counts = list(found_keywords.values())
            
            fig, ax = plt.subplots(figsize=(10, 5))
            bars = ax.barh(keywords, counts, color='purple')
            ax.set_xlabel('Frequency')
            ax.set_title('Keyword Frequency in Resume')
            
            # Add count labels to the bars
            for i, bar in enumerate(bars):
                ax.text(bar.get_width() + 0.1, bar.get_y() + bar.get_height()/2, 
                        str(counts[i]), va='center')
            
            st.pyplot(fig)
        
        # Display formatting recommendations
        st.subheader("ATS Format Recommendations")
        st.markdown("<div class='info-box'>", unsafe_allow_html=True)
        st.markdown(format_recommendations)
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Keyword insertion suggestions
        if missing_keywords:
            st.subheader("Keyword Integration Suggestions")
            
            integration_prompt = f"""
            Based on the resume content:
            {st.session_state.current_resume_text}
            
            Suggest natural ways to integrate these missing keywords:
            {', '.join(missing_keywords.keys())}
            
            Provide specific examples of how to modify or add content to include these keywords
            without making the resume sound forced or unnatural.
            """
            
            with st.spinner("Generating keyword suggestions..."):
                try:
                    integration_suggestions = ChatOpenAI(model_name="gpt-4", openai_api_key=OPENAI_API_KEY).predict(integration_prompt)
                    st.markdown("<div class='info-box'>", unsafe_allow_html=True)
                    st.markdown(integration_suggestions)
                    st.markdown("</div>", unsafe_allow_html=True)
                except Exception as e:
                    st.error(f"Error generating keyword suggestions: {e}")

# TAB 3: RESUME OPTIMIZER
with tab3:
    st.markdown("<h3 class='section-header'>Resume Optimization</h3>", unsafe_allow_html=True)
    
    if not st.session_state.current_resume_text or not st.session_state.current_job_description:
        st.warning("Please analyze a resume in the Resume Analysis tab first.")
    else:
        st.markdown("<div class='info-box'>", unsafe_allow_html=True)
        st.write("""
        Use AI to optimize specific sections of your resume to better match the job description.
        Choose which sections to improve, then generate an optimized version of your resume.
        """)
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Section optimization
        st.subheader("Optimize Resume Sections")
        
        col1, col2 = st.columns(2)
        
        with
